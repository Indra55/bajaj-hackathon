"""
Enhanced Document Processor with Domain Classification
Processes documents and classifies them by domain for better retrieval accuracy
"""

import base64
import requests
import re
from io import BytesIO
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass

try:
    from PyPDF2 import PdfReader
except ImportError:
    from pypdf import PdfReader
import docx

from .domain_classifier import DomainClassifier, Domain, DomainClassification
from .semantic_chunker import SemanticChunker, SemanticChunk

@dataclass
class ProcessedDocument:
    """Enhanced document with domain classification."""
    chunks: List[str]
    domain_classification: DomainClassification
    metadata: Dict
    filename: str
    total_length: int

class EnhancedDocumentProcessor:
    """Enhanced document processor with domain awareness."""
    
    def __init__(self, use_semantic_chunking: bool = True):
        self.min_chunk_size = 300  # Minimum words per chunk
        self.max_chunk_size = 600  # Maximum words per chunk
        self.overlap_words = 50    # Word overlap between chunks
        self.domain_classifier = DomainClassifier()
        self.use_semantic_chunking = use_semantic_chunking
        
        if use_semantic_chunking:
            self.semantic_chunker = SemanticChunker(
                min_sentences_per_chunk=3,
                max_sentences_per_chunk=12,
                similarity_threshold=0.3,
                max_chunk_words=500
            )
        
    def process_document_with_domain(self, document: dict) -> ProcessedDocument:
        """Process document with domain classification."""
        content = document.get('content')
        metadata = document.get('metadata', {})
        filename = metadata.get('filename', '')

        with self._get_content_stream(content) as stream:
            raw_text = self._extract_text(stream, filename)
            cleaned_text = self._clean_text(raw_text)
            
            # Classify document domain
            domain_classification = self.domain_classifier.classify_document(cleaned_text, filename)
            
            # Create chunks using semantic or traditional method
            if self.use_semantic_chunking:
                semantic_chunks = self.semantic_chunker.chunk_text(
                    cleaned_text, 
                    domain_classification.primary_domain.value
                )
                chunks = [chunk.content for chunk in semantic_chunks]
                print(f"Created {len(chunks)} semantic chunks")
            else:
                chunks = self._create_domain_aware_chunks(cleaned_text, domain_classification.primary_domain)
            
            # Enhanced metadata with domain info
            enhanced_metadata = {
                **metadata,
                'domain': domain_classification.primary_domain.value,
                'domain_confidence': domain_classification.confidence,
                'domain_keywords': domain_classification.keywords_found,
                'secondary_domains': [(d.value, conf) for d, conf in domain_classification.secondary_domains]
            }
            
            return ProcessedDocument(
                chunks=chunks,
                domain_classification=domain_classification,
                metadata=enhanced_metadata,
                filename=filename,
                total_length=len(cleaned_text)
            )

    def _get_content_stream(self, content: str) -> BytesIO:
        """Retrieves content as a stream from a URL or a Base64 string."""
        if content.startswith(('http://', 'https://')):
            response = requests.get(content, timeout=30, stream=True)
            response.raise_for_status()
            return BytesIO(response.content)
        else:
            decoded_content = base64.b64decode(content)
            return BytesIO(decoded_content)

    def _extract_text(self, content_stream: BytesIO, filename: str) -> str:
        """Extract text from document."""
        file_ext = filename.split('.')[-1].lower() if '.' in filename else ''

        text = ""
        if file_ext == 'pdf':
            reader = PdfReader(content_stream)
            for page in reader.pages:
                text += page.extract_text() or ""
        elif file_ext == 'docx':
            doc = docx.Document(content_stream)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            # Assume text content
            content_stream.seek(0)
            text = content_stream.read().decode('utf-8', errors='ignore')

        return text

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\"\'\/\%\$\@\#]', ' ', text)
        
        # Fix common OCR errors
        text = re.sub(r'\b(\w)\s+(\w)\b', r'\1\2', text)  # Fix spaced letters
        text = re.sub(r'(\d)\s+(\d)', r'\1\2', text)  # Fix spaced numbers
        
        # Normalize whitespace again
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text

    def _create_domain_aware_chunks(self, text: str, domain: Domain) -> List[str]:
        """Create chunks with domain-specific optimization."""
        words = text.split()
        
        if len(words) <= self.max_chunk_size:
            return [text]
        
        chunks = []
        current_chunk = []
        current_size = 0
        
        # Domain-specific sentence boundaries
        domain_boundaries = self._get_domain_boundaries(domain)
        
        sentences = self._split_into_sentences(text, domain_boundaries)
        
        for sentence in sentences:
            sentence_words = sentence.split()
            sentence_size = len(sentence_words)
            
            # If adding this sentence would exceed max size, finalize current chunk
            if current_size + sentence_size > self.max_chunk_size and current_chunk:
                chunk_text = ' '.join(current_chunk)
                if len(chunk_text.split()) >= self.min_chunk_size:
                    chunks.append(chunk_text)
                
                # Start new chunk with overlap
                overlap_words = current_chunk[-self.overlap_words:] if len(current_chunk) > self.overlap_words else current_chunk
                current_chunk = overlap_words + sentence_words
                current_size = len(current_chunk)
            else:
                current_chunk.extend(sentence_words)
                current_size += sentence_size
        
        # Add the last chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            if len(chunk_text.split()) >= self.min_chunk_size:
                chunks.append(chunk_text)
        
        return chunks

    def _get_domain_boundaries(self, domain: Domain) -> List[str]:
        """Get domain-specific sentence boundary patterns."""
        base_boundaries = [r'\.', r'\!', r'\?', r'\n\n']
        
        if domain == Domain.INSURANCE:
            # Insurance documents often have numbered clauses
            return base_boundaries + [r'\d+\.', r'Section\s+\d+', r'Clause\s+\d+', r'Article\s+\d+']
        
        elif domain == Domain.LEGAL:
            # Legal documents have articles, sections, subsections
            return base_boundaries + [r'Article\s+\d+', r'Section\s+\d+', r'Subsection\s+\d+', 
                                    r'\([a-z]\)', r'\([0-9]+\)', r'Paragraph\s+\d+']
        
        elif domain == Domain.HR:
            # HR documents have policy sections
            return base_boundaries + [r'Policy\s+\d+', r'Section\s+\d+', r'Procedure\s+\d+']
        
        elif domain == Domain.COMPLIANCE:
            # Compliance documents have requirements, controls
            return base_boundaries + [r'Requirement\s+\d+', r'Control\s+\d+', r'Standard\s+\d+']
        
        return base_boundaries

    def _split_into_sentences(self, text: str, boundaries: List[str]) -> List[str]:
        """Split text into sentences using domain-specific boundaries."""
        # Create a combined pattern
        pattern = '|'.join(f'({boundary})' for boundary in boundaries)
        
        # Split while keeping the delimiters
        parts = re.split(f'({pattern})', text)
        
        sentences = []
        current_sentence = ""
        
        for part in parts:
            if part and not re.match(f'^({pattern})$', part):
                current_sentence += part
            else:
                if current_sentence.strip():
                    sentences.append(current_sentence.strip())
                    current_sentence = ""
        
        if current_sentence.strip():
            sentences.append(current_sentence.strip())
        
        # Filter out very short sentences
        sentences = [s for s in sentences if len(s.split()) >= 5]
        
        return sentences

    def get_domain_specific_chunks(self, processed_doc: ProcessedDocument, 
                                 query_domain: Domain) -> List[str]:
        """Get chunks that are most relevant to the query domain."""
        if query_domain == Domain.GENERAL or query_domain == processed_doc.domain_classification.primary_domain:
            return processed_doc.chunks
        
        # If domains don't match exactly, check secondary domains
        secondary_domain_values = [d.value for d, _ in processed_doc.domain_classification.secondary_domains]
        if query_domain.value in secondary_domain_values:
            return processed_doc.chunks
        
        # If no domain match, return empty list (don't search this document)
        return []

    def create_document_index(self, processed_docs: List[ProcessedDocument]) -> Dict[Domain, List[int]]:
        """Create an index of documents by domain."""
        domain_index = {domain: [] for domain in Domain}
        
        for i, doc in enumerate(processed_docs):
            # Add to primary domain
            domain_index[doc.domain_classification.primary_domain].append(i)
            
            # Add to secondary domains if confidence is high enough
            for domain, confidence in doc.domain_classification.secondary_domains:
                if confidence > 0.2:  # Threshold for secondary domain inclusion
                    domain_index[domain].append(i)
        
        return domain_index
